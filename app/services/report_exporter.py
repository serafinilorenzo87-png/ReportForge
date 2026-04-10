from __future__ import annotations

from collections import Counter
from pathlib import Path
from tempfile import NamedTemporaryFile
import re

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.models.finding import Finding
from app.models.project import Project
from app.models.reference import Reference


SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Info"]
ACCENT_RGB = RGBColor(37, 99, 235)
ACCENT_HEX = "#2563eb"
LIGHT_BLUE_HEX = "#EAF2FF"
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/*?:"<>|]', "", value).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "reportforge_project"


def build_severity_summary(findings: list[Finding]) -> dict[str, int]:
    counts = Counter((finding.severity or "Info").strip().capitalize() for finding in findings)
    return {severity: counts.get(severity, 0) for severity in SEVERITY_ORDER}


def _safe_text(value: str | None, fallback: str = "Not specified") -> str:
    text = (value or "").strip()
    return text if text else fallback


def _add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    size: int | None = None,
):
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = ACCENT_RGB

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def _add_cover_page(doc: Document, project: Project) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "ReportForge", bold=True, color=ACCENT_RGB, size=24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Penetration Test Report", bold=True, size=22)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, _safe_text(project.project_name), bold=True, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        p,
        f"{_safe_text(project.client_name)}  •  {_safe_text(project.target_name)}  •  {_safe_text(project.engagement_type)}",
        size=11,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        p,
        f"Assessment Window: {_safe_text(project.assessment_start)} to {_safe_text(project.assessment_end)}",
        size=11,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        p,
        f"Risk Rating: {_safe_text(project.risk_rating)}",
        bold=True,
        color=ACCENT_RGB,
        size=12,
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Generated With ReportForge", size=10)

    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_key_value_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
        _set_cell_shading(row[0], "EAF2FF")

    doc.add_paragraph()


def _generate_severity_chart(findings: list[Finding]) -> Path | None:
    summary = build_severity_summary(findings)
    labels = [severity for severity in SEVERITY_ORDER if summary[severity] > 0]
    values = [summary[severity] for severity in labels]

    if not values:
        return None

    fig, ax = plt.subplots(figsize=(6.5, 4.2))
    ax.pie(values, labels=labels, autopct="%1.0f%%", startangle=140)
    ax.set_title("Severity Distribution")

    temp_file = NamedTemporaryFile(delete=False, suffix=".png")
    temp_path = Path(temp_file.name)
    temp_file.close()

    fig.savefig(temp_path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return temp_path


def _add_summary_section(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(_safe_text(text, "No summary provided."))


def _build_short_term_actions(findings: list[Finding]) -> str:
    critical_high = [
        f.title
        for f in findings
        if (f.severity or "").strip().capitalize() in {"Critical", "High"}
    ]
    if not critical_high:
        return (
            "No urgent short-term actions identified. Prioritize validation of medium-priority issues, "
            "basic hardening, and routine security hygiene improvements."
        )

    joined = ", ".join(critical_high[:6])
    return (
        "Immediate remediation should focus on the highest-risk findings identified during the assessment, "
        f"including: {joined}. Exposure reduction, access restriction, patching, and configuration hardening "
        "should be treated as the first remediation wave."
    )


def _build_long_term_actions(project: Project, findings: list[Finding]) -> str:
    return (
        "Long-term improvement should include repeatable hardening processes, stronger credential policies, "
        "periodic validation of exposed assets, continuous vulnerability management, and clearer ownership for "
        f"the environment described as {_safe_text(project.environment_name, 'the assessed environment')}. "
        "Future assessments should also verify whether corrective actions have been fully implemented and sustained over time."
    )


def _is_image_file(path_str: str) -> bool:
    try:
        return Path(path_str).suffix.lower() in IMAGE_EXTENSIONS
    except Exception:
        return False


def _existing_image_paths(attachments: list[str]) -> list[Path]:
    results: list[Path] = []
    for attachment in attachments:
        if not _is_image_file(attachment):
            continue
        path = Path(attachment)
        if path.exists() and path.is_file():
            results.append(path)
    return results


def _existing_non_image_paths(attachments: list[str]) -> list[str]:
    results: list[str] = []
    for attachment in attachments:
        if _is_image_file(attachment):
            continue
        path = Path(attachment)
        if path.exists() and path.is_file():
            results.append(str(path))
        else:
            results.append(attachment)
    return results


def _fit_image_dimensions(
    image_path: Path,
    max_width_in: float,
    max_height_in: float,
) -> tuple[float, float]:
    with PILImage.open(image_path) as img:
        width_px, height_px = img.size

    if width_px <= 0 or height_px <= 0:
        return max_width_in, min(max_height_in, max_width_in)

    width_in = width_px / 96
    height_in = height_px / 96
    ratio = min(max_width_in / width_in, max_height_in / height_in, 1.0)
    return width_in * ratio, height_in * ratio


def _add_docx_attachment_images(doc: Document, attachments: list[str]) -> None:
    image_paths = _existing_image_paths(attachments)
    other_paths = _existing_non_image_paths(attachments)

    if image_paths:
        doc.add_heading("Evidence Screenshots", level=3)
        for image_path in image_paths:
            doc.add_paragraph(image_path.name)
            width_in, height_in = _fit_image_dimensions(image_path, 6.2, 5.8)
            doc.add_picture(
                str(image_path),
                width=Inches(width_in),
                height=Inches(height_in),
            )
            doc.add_paragraph()

    if other_paths:
        doc.add_heading("Other Attachments", level=3)
        for attachment in other_paths:
            doc.add_paragraph(attachment)


def _add_pdf_attachment_images(
    story: list,
    attachments: list[str],
    body_style,
    heading_style,
) -> None:
    image_paths = _existing_image_paths(attachments)
    other_paths = _existing_non_image_paths(attachments)

    if image_paths:
        story.append(Paragraph("Evidence Screenshots", heading_style))
        for image_path in image_paths:
            story.append(Paragraph(image_path.name, body_style))
            width_in, height_in = _fit_image_dimensions(image_path, 5.8, 5.4)
            story.append(
                Image(
                    str(image_path),
                    width=width_in * 72.0,
                    height=height_in * 72.0,
                )
            )
            story.append(Spacer(1, 0.2 * cm))

    if other_paths:
        story.append(Paragraph("Other Attachments", heading_style))
        for attachment in other_paths:
            story.append(Paragraph(f"• {attachment}", body_style))


def build_markdown_report(
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> str:
    severity_summary = build_severity_summary(findings)
    total_references = sum(len(refs) for refs in references_by_finding.values())

    lines: list[str] = []
    lines.append(f"# {project.project_name}")
    lines.append("")
    lines.append("## Engagement Overview")
    lines.append("")
    lines.append(f"- **Client:** {_safe_text(project.client_name)}")
    lines.append(f"- **Target:** {_safe_text(project.target_name)}")
    lines.append(f"- **Platform:** {_safe_text(project.platform)}")
    lines.append(f"- **Engagement Type:** {_safe_text(project.engagement_type)}")
    lines.append(
        f"- **Assessment Window:** {_safe_text(project.assessment_start)} to {_safe_text(project.assessment_end)}"
    )
    lines.append(f"- **Risk Rating:** {_safe_text(project.risk_rating)}")
    lines.append("")

    sections = [
        ("Stakeholder Summary", project.stakeholder_summary),
        ("Executive Summary", project.executive_summary),
        ("Technical Summary", project.technical_summary),
        ("Scope Summary", project.scope_summary),
        ("Out of Scope", project.out_of_scope),
        ("Attack Surface", project.attack_surface),
        (
            "Targets and Assets",
            f"IPs: {_safe_text(project.target_ips)}\n\nDomains: {_safe_text(project.target_domains)}",
        ),
        (
            "Applications and Environment",
            "\n".join(
                [
                    f"Web Application: {_safe_text(project.web_app_name)}",
                    f"Mobile Application: {_safe_text(project.mobile_app_name)}",
                    f"Internal Application: {_safe_text(project.internal_app_name)}",
                    f"Environment: {_safe_text(project.environment_name)}",
                ]
            ),
        ),
        ("Methodology", project.methodology_summary),
        ("Standards Used", project.standards_used),
    ]

    for title, text in sections:
        lines.append(f"## {title}")
        lines.append("")
        lines.append(_safe_text(text, "Not provided."))
        lines.append("")

    lines.append("## Severity Summary")
    lines.append("")
    lines.append(
        f"This assessment contains **{len(findings)}** documented finding(s) "
        f"and **{total_references}** linked reference(s)."
    )
    lines.append("")
    for severity in SEVERITY_ORDER:
        lines.append(f"- **{severity}:** {severity_summary[severity]}")
    lines.append("")

    lines.append("## Detailed Findings")
    lines.append("")

    if findings:
        for index, finding in enumerate(findings, start=1):
            lines.append(f"### {index}. {finding.title}")
            lines.append("")
            lines.append(f"- **Severity:** {finding.severity}")
            lines.append("")
            lines.append("#### Description")
            lines.append("")
            lines.append(_safe_text(finding.description, "No description provided."))
            lines.append("")
            lines.append("#### Evidence")
            lines.append("")
            lines.append(_safe_text(finding.evidence, "No evidence provided."))
            lines.append("")

            attachments = finding.get_attachments()
            if attachments:
                lines.append("#### Evidence Attachments")
                lines.append("")
                for attachment in attachments:
                    lines.append(f"- `{attachment}`")
                lines.append("")

            lines.append("#### Remediation")
            lines.append("")
            lines.append(_safe_text(finding.remediation, "No remediation provided."))
            lines.append("")

            finding_references = references_by_finding.get(finding.id, [])
            if finding_references:
                lines.append("#### References")
                lines.append("")
                for reference in finding_references:
                    ref_line = f"- **{reference.title}** ({reference.reference_type})"
                    if (reference.url or "").strip():
                        ref_line += f" — {reference.url.strip()}"
                    lines.append(ref_line)
                    if (reference.note or "").strip():
                        lines.append(f"  - Note: {reference.note.strip()}")
                lines.append("")
    else:
        lines.append("No findings documented.")
        lines.append("")

    lines.append("## Short-Term Actions")
    lines.append("")
    lines.append(_build_short_term_actions(findings))
    lines.append("")
    lines.append("## Long-Term Improvements")
    lines.append("")
    lines.append(_build_long_term_actions(project, findings))
    lines.append("")
    lines.append("## Conclusions")
    lines.append("")
    lines.append(_safe_text(project.conclusions, "No final conclusion provided."))
    lines.append("")
    lines.append("## Generated With")
    lines.append("")
    lines.append("This report was created with ReportForge.")
    lines.append("")

    return "\n".join(lines)


def build_text_report(
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> str:
    severity_summary = build_severity_summary(findings)
    total_references = sum(len(refs) for refs in references_by_finding.values())

    lines: list[str] = []
    lines.append(project.project_name)
    lines.append("=" * len(project.project_name))
    lines.append("")
    lines.append("ENGAGEMENT OVERVIEW")
    lines.append("-------------------")
    lines.append(f"Client: {_safe_text(project.client_name)}")
    lines.append(f"Target: {_safe_text(project.target_name)}")
    lines.append(f"Platform: {_safe_text(project.platform)}")
    lines.append(f"Engagement Type: {_safe_text(project.engagement_type)}")
    lines.append(
        f"Assessment Window: {_safe_text(project.assessment_start)} to {_safe_text(project.assessment_end)}"
    )
    lines.append(f"Risk Rating: {_safe_text(project.risk_rating)}")
    lines.append("")

    lines.append("STAKEHOLDER SUMMARY")
    lines.append("-------------------")
    lines.append(_safe_text(project.stakeholder_summary, "No stakeholder summary provided."))
    lines.append("")
    lines.append("EXECUTIVE SUMMARY")
    lines.append("-----------------")
    lines.append(_safe_text(project.executive_summary, "No executive summary provided."))
    lines.append("")
    lines.append("TECHNICAL SUMMARY")
    lines.append("-----------------")
    lines.append(_safe_text(project.technical_summary, "No technical summary provided."))
    lines.append("")

    lines.append("SEVERITY SUMMARY")
    lines.append("----------------")
    lines.append(
        f"This assessment contains {len(findings)} documented finding(s) and {total_references} linked reference(s)."
    )
    for severity in SEVERITY_ORDER:
        lines.append(f"{severity}: {severity_summary[severity]}")
    lines.append("")

    lines.append("DETAILED FINDINGS")
    lines.append("-----------------")
    if findings:
        for index, finding in enumerate(findings, start=1):
            lines.append(f"{index}. {finding.title}")
            lines.append(f"Severity: {finding.severity}")
            lines.append("")
            lines.append("Description:")
            lines.append(_safe_text(finding.description, "No description provided."))
            lines.append("")
            lines.append("Evidence:")
            lines.append(_safe_text(finding.evidence, "No evidence provided."))
            lines.append("")

            attachments = finding.get_attachments()
            if attachments:
                lines.append("Evidence Attachments:")
                for attachment in attachments:
                    lines.append(f"- {attachment}")
                lines.append("")

            lines.append("Remediation:")
            lines.append(_safe_text(finding.remediation, "No remediation provided."))
            lines.append("")

            finding_references = references_by_finding.get(finding.id, [])
            if finding_references:
                lines.append("References:")
                for reference in finding_references:
                    ref_line = f"- {reference.title} ({reference.reference_type})"
                    if (reference.url or "").strip():
                        ref_line += f" - {reference.url.strip()}"
                    lines.append(ref_line)
                    if (reference.note or "").strip():
                        lines.append(f"  Note: {reference.note.strip()}")
                lines.append("")

            lines.append("-" * 60)
            lines.append("")
    else:
        lines.append("No findings documented.")
        lines.append("")

    lines.append("SHORT-TERM ACTIONS")
    lines.append("------------------")
    lines.append(_build_short_term_actions(findings))
    lines.append("")
    lines.append("LONG-TERM IMPROVEMENTS")
    lines.append("----------------------")
    lines.append(_build_long_term_actions(project, findings))
    lines.append("")
    lines.append("CONCLUSIONS")
    lines.append("-----------")
    lines.append(_safe_text(project.conclusions, "No final conclusion provided."))
    lines.append("")
    lines.append("GENERATED WITH")
    lines.append("--------------")
    lines.append("This report was created with ReportForge.")
    lines.append("")

    return "\n".join(lines)


def export_markdown_report(
    output_dir: Path,
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = sanitize_filename(project.project_name) + ".md"
    output_path = output_dir / filename
    output_path.write_text(
        build_markdown_report(project, findings, references_by_finding),
        encoding="utf-8",
    )
    return output_path


def export_text_report(
    output_dir: Path,
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = sanitize_filename(project.project_name) + ".txt"
    output_path = output_dir / filename
    output_path.write_text(
        build_text_report(project, findings, references_by_finding),
        encoding="utf-8",
    )
    return output_path


def export_docx_report(
    output_dir: Path,
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _style_doc(doc)
    _add_cover_page(doc, project)

    _add_key_value_table(
        doc,
        "Engagement Overview",
        [
            ("Project Name", _safe_text(project.project_name)),
            ("Client Name", _safe_text(project.client_name)),
            ("Target Name", _safe_text(project.target_name)),
            ("Platform", _safe_text(project.platform)),
            ("Engagement Type", _safe_text(project.engagement_type)),
            ("Assessment Start", _safe_text(project.assessment_start)),
            ("Assessment End", _safe_text(project.assessment_end)),
            ("Risk Rating", _safe_text(project.risk_rating)),
            ("Environment", _safe_text(project.environment_name)),
            ("Web App", _safe_text(project.web_app_name)),
            ("Mobile App", _safe_text(project.mobile_app_name)),
            ("Internal App", _safe_text(project.internal_app_name)),
            ("Target IPs", _safe_text(project.target_ips)),
            ("Target Domains", _safe_text(project.target_domains)),
        ],
    )

    _add_summary_section(doc, "Scope Summary", project.scope_summary)
    _add_summary_section(doc, "Out of Scope", project.out_of_scope)
    _add_summary_section(doc, "Attack Surface", project.attack_surface)
    _add_summary_section(doc, "Methodology", project.methodology_summary)
    _add_summary_section(doc, "Standards Used", project.standards_used)
    _add_summary_section(doc, "Stakeholder Summary", project.stakeholder_summary)
    _add_summary_section(doc, "Executive Summary", project.executive_summary)
    _add_summary_section(doc, "Technical Summary", project.technical_summary)

    summary = build_severity_summary(findings)
    chart_path = None

    try:
        doc.add_heading("Severity Summary", level=1)
        severity_table = doc.add_table(rows=1, cols=2)
        severity_table.style = "Table Grid"
        severity_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = severity_table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"
        _set_cell_shading(hdr[0], "EAF2FF")
        _set_cell_shading(hdr[1], "EAF2FF")

        for severity in SEVERITY_ORDER:
            row = severity_table.add_row().cells
            row[0].text = severity
            row[1].text = str(summary[severity])

        doc.add_paragraph()

        chart_path = _generate_severity_chart(findings)
        if chart_path is not None and chart_path.exists():
            doc.add_picture(str(chart_path), width=Inches(5.8))
            doc.add_paragraph()

        doc.add_heading("Detailed Findings", level=1)

        if findings:
            for index, finding in enumerate(findings, start=1):
                doc.add_heading(f"{index}. {finding.title}", level=2)

                p = doc.add_paragraph()
                _add_run(p, "Severity: ", bold=True)
                _add_run(p, _safe_text(finding.severity, "Info"))

                doc.add_heading("Description", level=3)
                doc.add_paragraph(_safe_text(finding.description, "No description provided."))

                doc.add_heading("Evidence", level=3)
                doc.add_paragraph(_safe_text(finding.evidence, "No evidence provided."))

                attachments = finding.get_attachments()
                if attachments:
                    _add_docx_attachment_images(doc, attachments)

                doc.add_heading("Remediation", level=3)
                doc.add_paragraph(_safe_text(finding.remediation, "No remediation provided."))

                finding_references = references_by_finding.get(finding.id, [])
                if finding_references:
                    doc.add_heading("References", level=3)
                    for reference in finding_references:
                        text = f"{reference.reference_type}: {reference.title}"
                        if (reference.url or "").strip():
                            text += f" — {reference.url.strip()}"
                        doc.add_paragraph(text)
                        if (reference.note or "").strip():
                            doc.add_paragraph(f"Note: {reference.note.strip()}")

                doc.add_paragraph()
        else:
            doc.add_paragraph("No findings documented.")

        doc.add_heading("Short-Term Actions", level=1)
        doc.add_paragraph(_build_short_term_actions(findings))

        doc.add_heading("Long-Term Improvements", level=1)
        doc.add_paragraph(_build_long_term_actions(project, findings))

        doc.add_heading("Conclusions", level=1)
        doc.add_paragraph(_safe_text(project.conclusions, "No final conclusion provided."))

        doc.add_heading("Generated With", level=1)
        doc.add_paragraph("This report was created with ReportForge.")
    finally:
        if chart_path is not None and chart_path.exists():
            chart_path.unlink(missing_ok=True)

    filename = sanitize_filename(project.project_name) + ".docx"
    output_path = output_dir / filename
    doc.save(output_path)
    return output_path


def export_pdf_report(
    output_dir: Path,
    project: Project,
    findings: list[Finding],
    references_by_finding: dict[int, list[Reference]],
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = sanitize_filename(project.project_name) + ".pdf"
    output_path = output_dir / filename

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "RFTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        textColor=colors.HexColor(ACCENT_HEX),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "RFSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        alignment=TA_CENTER,
        textColor=colors.black,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "RFHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        textColor=colors.HexColor(ACCENT_HEX),
        spaceAfter=8,
        spaceBefore=10,
    )
    subheading_style = ParagraphStyle(
        "RFSubHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        textColor=colors.HexColor(ACCENT_HEX),
        spaceAfter=5,
        spaceBefore=6,
    )
    body_style = ParagraphStyle(
        "RFBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.6,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=6,
    )

    story = []
    story.append(Spacer(1, 1.2 * cm))
    story.append(Paragraph(_safe_text(project.project_name), title_style))
    story.append(Paragraph("Penetration Test Report", subtitle_style))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(f"Client  {_safe_text(project.client_name)}", subtitle_style))
    story.append(Paragraph(f"Target  {_safe_text(project.target_name)}", subtitle_style))
    story.append(Paragraph(f"Engagement Type  {_safe_text(project.engagement_type)}", subtitle_style))
    story.append(
        Paragraph(
            f"Assessment Window  {_safe_text(project.assessment_start)} to {_safe_text(project.assessment_end)}",
            subtitle_style,
        )
    )
    story.append(Paragraph(f"Risk Rating  {_safe_text(project.risk_rating)}", subtitle_style))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Generated With ReportForge", subtitle_style))
    story.append(PageBreak())

    story.append(Paragraph("Engagement Overview", heading_style))
    overview_rows = [
        ["Project Name", _safe_text(project.project_name)],
        ["Client Name", _safe_text(project.client_name)],
        ["Target Name", _safe_text(project.target_name)],
        ["Platform", _safe_text(project.platform)],
        ["Engagement Type", _safe_text(project.engagement_type)],
        ["Assessment Start", _safe_text(project.assessment_start)],
        ["Assessment End", _safe_text(project.assessment_end)],
        ["Risk Rating", _safe_text(project.risk_rating)],
        ["Environment", _safe_text(project.environment_name)],
        ["Web App", _safe_text(project.web_app_name)],
        ["Mobile App", _safe_text(project.mobile_app_name)],
        ["Internal App", _safe_text(project.internal_app_name)],
        ["Target IPs", _safe_text(project.target_ips)],
        ["Target Domains", _safe_text(project.target_domains)],
    ]
    overview_table = Table(overview_rows, colWidths=[4.2 * cm, 11.8 * cm])
    overview_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(LIGHT_BLUE_HEX)),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.2),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(overview_table)
    story.append(Spacer(1, 0.25 * cm))

    section_pairs = [
        ("Scope Summary", project.scope_summary),
        ("Out of Scope", project.out_of_scope),
        ("Attack Surface", project.attack_surface),
        ("Methodology", project.methodology_summary),
        ("Standards Used", project.standards_used),
        ("Stakeholder Summary", project.stakeholder_summary),
        ("Executive Summary", project.executive_summary),
        ("Technical Summary", project.technical_summary),
    ]

    for title, text in section_pairs:
        story.append(Paragraph(title, heading_style))
        story.append(Paragraph(_safe_text(text, "No summary provided."), body_style))

    summary = build_severity_summary(findings)
    story.append(Paragraph("Severity Summary", heading_style))

    severity_rows = [["Severity", "Count"]]
    for severity in SEVERITY_ORDER:
        severity_rows.append([severity, str(summary[severity])])

    severity_table = Table(severity_rows, colWidths=[7 * cm, 3 * cm])
    severity_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(LIGHT_BLUE_HEX)),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (1, 1), (1, -1), "CENTER"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(severity_table)
    story.append(Spacer(1, 0.25 * cm))

    chart_path = _generate_severity_chart(findings)
    try:
        if chart_path is not None and chart_path.exists():
            story.append(Image(str(chart_path), width=14.0 * cm, height=9.0 * cm))
            story.append(Spacer(1, 0.25 * cm))

        story.append(Paragraph("Detailed Findings", heading_style))

        if findings:
            for index, finding in enumerate(findings, start=1):
                story.append(Paragraph(f"{index}. {finding.title}", heading_style))
                story.append(
                    Paragraph(
                        f"<b>Severity:</b> {_safe_text(finding.severity, 'Info')}",
                        body_style,
                    )
                )

                story.append(Paragraph("Description", subheading_style))
                story.append(Paragraph(_safe_text(finding.description, "No description provided."), body_style))

                story.append(Paragraph("Evidence", subheading_style))
                story.append(Paragraph(_safe_text(finding.evidence, "No evidence provided."), body_style))

                attachments = finding.get_attachments()
                if attachments:
                    _add_pdf_attachment_images(story, attachments, body_style, subheading_style)

                story.append(Paragraph("Remediation", subheading_style))
                story.append(Paragraph(_safe_text(finding.remediation, "No remediation provided."), body_style))

                finding_references = references_by_finding.get(finding.id, [])
                if finding_references:
                    story.append(Paragraph("References", subheading_style))
                    for reference in finding_references:
                        line = f"{reference.reference_type}: {reference.title}"
                        if (reference.url or "").strip():
                            line += f" — {reference.url.strip()}"
                        story.append(Paragraph(line, body_style))
                        if (reference.note or "").strip():
                            story.append(Paragraph(f"Note: {reference.note.strip()}", body_style))

                story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Paragraph("No findings documented.", body_style))

        story.append(Paragraph("Short-Term Actions", heading_style))
        story.append(Paragraph(_build_short_term_actions(findings), body_style))

        story.append(Paragraph("Long-Term Improvements", heading_style))
        story.append(Paragraph(_build_long_term_actions(project, findings), body_style))

        story.append(Paragraph("Conclusions", heading_style))
        story.append(Paragraph(_safe_text(project.conclusions, "No final conclusion provided."), body_style))

        story.append(Paragraph("Generated With", heading_style))
        story.append(Paragraph("This report was created with ReportForge.", body_style))

        doc.build(story)
    finally:
        if chart_path is not None and chart_path.exists():
            chart_path.unlink(missing_ok=True)

    return output_path